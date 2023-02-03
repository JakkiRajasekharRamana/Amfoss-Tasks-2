import docx
doc = docx.Document('invitation.docx')
name = open('guests.txt')
name = name.read()
name = name.split('/n')
for i in name:
    doc.add_paragraph('It would be a pleasure to have the company of').style = 'outline'
    doc.add_paragraph(name).style = 'Name'
    doc.add_paragraph('at 11010 Memory Lane on the Evening of').style = 'imprint'
    doc.add_paragraph('April 1st').style = 'all_caps'
    date_line = doc.add_paragraph('at 7 o\'clock')
    date_line.style = 'shadow'
